"""DocumentExtractor implementation backed by the Anthropic API.

One adapter, four input shapes:
- plain text  -> sent as text
- image       -> sent as a base64 image block (vision)
- PDF         -> sent as a base64 document block (native PDF support)
- docx        -> text extracted locally with python-docx, then sent as text

The model is asked for strict JSON; we parse defensively.
API reference: https://docs.claude.com/en/api/overview
"""
from __future__ import annotations

import base64
import io
import json
import re

import anthropic

from app.domain.models import ExtractedBook, IncomingMessage, MediaKind

_SYSTEM = (
    "You extract book lists from Nigerian school booklists sent to a bookstore. "
    "The input may be a photo of a printed or handwritten list, a PDF, or plain text. "
    "Return ONLY a JSON array, no prose, no markdown fences. Each element: "
    '{"title": str, "author_or_publisher": str|null, "grade": str|null, "quantity": int}. '
    "Rules: keep titles as written (do not invent authors); grade is the class the list "
    "is for (e.g. 'Primary 4', 'JSS 2', 'SS 1') if stated anywhere on the document, "
    "applied to every book on that list; quantity defaults to 1 unless the list says "
    "otherwise; ignore non-book items (uniforms, fees, stationery) unless they are "
    "clearly books/workbooks. If the document contains no book list, return []."
)

_FENCE = re.compile(r"^```(?:json)?|```$", re.MULTILINE)


class ClaudeExtractor:
    def __init__(self, api_key: str, model: str):
        self._client = anthropic.Anthropic(api_key=api_key)
        self._model = model

    # ---- DocumentExtractor port ----

    def extract(self, message: IncomingMessage) -> list[ExtractedBook]:
        content = self._build_content(message)
        response = self._client.messages.create(
            model=self._model,
            max_tokens=4000,
            system=_SYSTEM,
            messages=[{"role": "user", "content": content}],
        )
        text = "".join(
            block.text for block in response.content if block.type == "text"
        )
        return self._parse(text)

    # ---- internals ----

    def _build_content(self, msg: IncomingMessage) -> list[dict]:
        instruction = {
            "type": "text",
            "text": "Extract the book list from this message."
            + (f" Caption/text from sender: {msg.text}" if msg.text else ""),
        }
        if msg.kind == MediaKind.TEXT:
            return [
                {
                    "type": "text",
                    "text": f"Extract the book list from this text:\n\n{msg.text}",
                }
            ]
        if msg.kind == MediaKind.IMAGE:
            return [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": msg.media_mime or "image/jpeg",
                        "data": base64.b64encode(msg.media_bytes or b"").decode(),
                    },
                },
                instruction,
            ]
        if msg.kind == MediaKind.PDF:
            return [
                {
                    "type": "document",
                    "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": base64.b64encode(msg.media_bytes or b"").decode(),
                    },
                },
                instruction,
            ]
        if msg.kind == MediaKind.DOCX:
            text = self._docx_to_text(msg.media_bytes or b"")
            return [
                {
                    "type": "text",
                    "text": f"Extract the book list from this document text:\n\n{text}",
                }
            ]
        raise ValueError(f"Unsupported media kind: {msg.kind}")

    @staticmethod
    def _docx_to_text(data: bytes) -> str:
        from docx import Document  # python-docx; imported lazily

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)

    @staticmethod
    def _parse(raw: str) -> list[ExtractedBook]:
        cleaned = _FENCE.sub("", raw).strip()
        # Grab the outermost JSON array even if the model added stray text.
        start, end = cleaned.find("["), cleaned.rfind("]")
        if start == -1 or end == -1:
            raise ValueError(f"Extractor returned non-JSON output: {raw[:200]}")
        data = json.loads(cleaned[start : end + 1])
        books: list[ExtractedBook] = []
        for entry in data:
            title = str(entry.get("title", "")).strip()
            if not title:
                continue
            qty_raw = entry.get("quantity", 1)
            try:
                qty = max(1, int(qty_raw))
            except (TypeError, ValueError):
                qty = 1
            books.append(
                ExtractedBook(
                    title=title,
                    author_or_publisher=(entry.get("author_or_publisher") or None),
                    grade=(entry.get("grade") or None),
                    quantity=qty,
                )
            )
        return books
